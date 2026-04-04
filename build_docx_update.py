#!/usr/bin/env python3
"""
在中文版 interim report 中插入两个新的文献综述小节：
  2.3 AML 系统的可解释性与 LLM 解释层
  2.4 隐私保护合规：零知识证明与 Privacy Pools
以及追加新的参考文献 [17]-[26]。
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

INPUT  = "/Users/tangliam/CriptoAnalyst/1st Interim Report - 中文版.docx"
OUTPUT = "/Users/tangliam/CriptoAnalyst/1st Interim Report - 中文版 (updated).docx"

doc = Document(INPUT)

# ========== 1. 找到插入点 ==========
insert_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "3. 系统建模与结构" and p.style.name == "Heading 1":
        insert_idx = i
        break

if insert_idx is None:
    raise RuntimeError("找不到 '3. 系统建模与结构' 标题")

print(f"将在 paragraph index {insert_idx} 前插入新内容")

# ========== 2. 准备新段落内容 ==========
new_paragraphs = []

def add(style, text):
    new_paragraphs.append((style, text))

# --- 2.3 AML 系统的可解释性与 LLM 解释层 ---
add("Heading 2", "2.3 AML 系统的可解释性与 LLM 解释层")

add("Heading 3", "2.3.1 可解释性的监管刚需")

add("Normal",
    "AML 检测系统的可解释性并非\u201c锦上添花\u201d，而是监管的硬性要求。"
    "FATF Recommendation 20 要求金融机构在提交可疑交易报告（Suspicious Transaction Report, STR）时必须包含"为什么认为该交易可疑"的文字说明。"
    "美国《银行保密法》（Bank Secrecy Act）同样要求可疑活动报告（SAR）以自然语言描述可疑行为的具体模式。"
    "欧盟《人工智能法案》（EU AI Act, Regulation 2024/1689）于 2024 年正式生效，将用于 AML/CFT 合规的 AI 系统归类为"高风险 AI 系统"（Annex III）[17]，"
    "要求此类系统必须满足决策逻辑的充分透明性(transparency)、人类监督机制(human oversight)以及可向监管机构和终端用户提供完整的决策解释(explainability)。"
    "违规者最高面临 3,500 万欧元或全球营业额 7% 的罚款。")

add("Normal",
    "这产生了一个核心矛盾：检测效果最好的模型（GNN、深度学习）恰恰是最不可解释的；而最可解释的方法（规则引擎）检测能力有限。"
    "LLM 的价值正好在这里——它不替代检测，而是把检测结果翻译成人类能理解的语言。")

add("Heading 3", "2.3.2 XAI 在 AML 中的应用")

add("Normal",
    "近年来，将可解释性工具（SHAP、LIME 等）引入 AML 模型的研究迅速增长。"
    "Kute 等人（2026）提出了一个端到端的可复现 SHAP 框架，用于解释 AML 模型的决策过程，强调公平性与监管合规的结合 [18]。"
    "该工作系统地展示了如何将特征重要性归因转化为合规部门可理解的报告。")

add("Normal",
    "在区块链场景下，Watson, Richards & Schiff（2025）提出了一个具有代表性的三层架构 [19]："
    "第一层为 GNN 检测层（GCN-GRU 混合模型，准确率 0.9470，AUC-ROC 0.9807）；"
    "第二层为 GraphLIME 归因层（识别哪些特征驱动了分类结果）；"
    "第三层为 LLM 解释层（将归因结果转化为自然语言叙述）。"
    "该架构在 Elliptic++ 数据集上验证，代表了 GNN + XAI + LLM 三层融合的最新范式。")

add("Heading 3", "2.3.3 LLM 在链上交易解释中的独立应用")

add("Normal",
    "除了作为 GNN 的解释层，LLM 也被直接应用于链上交易数据的分析和解释。"
    "Nicholls 等人（2024）在 Large Language Model XAI Approach for Illicit Activity Investigation in Bitcoin"
    "（发表于 Springer Neural Computing and Applications，IF 4.7）中 [20]，"
    "展示了一种不依赖传统 ML 检测器的方法：直接将 Bitcoin 交易数据输入 LLM 生成自然语言叙述，"
    "再提取叙述的嵌入向量计算相似度，从而发现其他非法交易。"
    "该方法的意义在于证明了 LLM 对链上交易数据能产生有用的、可操作的解释，且这些解释本身可以作为一种特征用于下游检测。")

add("Normal",
    "Sun 等人（2024）发表了第一篇系统综述 LLM 在区块链安全领域应用的论文 "
    "Large Language Models for Blockchain Security: A Systematic Literature Review [21]，"
    "覆盖异常检测、智能合约审计、交易分析等方向，为该领域的研究全景提供了结构化梳理。")

add("Heading 3", "2.3.4 与本研究的关系")

add("Normal",
    "本研究的规则引擎（BFS 追踪 + 风险评分）在可解释性方面具有天然优势："
    "每一步决策（为什么该节点被标为 suspect、风险路径是什么、衰减是怎么算的）都有完整的因果链条，不存在"黑盒"问题。"
    "这使得引入 LLM 作为解释层的路径更为直接——不需要 SHAP/LIME 等 post-hoc 归因工具来"反向推测"模型决策，"
    "而是直接将已有的结构化追踪结果（JSON）转化为自然语言风险报告，天然满足 SAR/STR 的叙述性要求。")

# --- 2.4 隐私保护合规：零知识证明与 Privacy Pools ---
add("Heading 2", "2.4 隐私保护合规：零知识证明与 Privacy Pools")

add("Heading 3", "2.4.1 隐私与合规的核心矛盾")

add("Normal",
    "区块链 AML 领域存在一个根本性张力：隐私保护（Privacy）与合规审查（Compliance）长期被视为不可兼得的两极。"
    "用户有合法的隐私需求（不暴露资产状况和交易细节），但合规框架要求资金来源的可追溯性。")

add("Normal",
    "2022 年 8 月，美国 OFAC 首次制裁了 Tornado Cash 智能合约——这是历史上首次对一个开源、不可变的代码（而非人或组织）实施制裁，引发了学术界和法律界的广泛讨论。"
    "Brownworth, Durfee, Lee & Martin（2024）在纽约联储的工作论文中对此进行了系统的实证分析 [22]："
    "制裁公告后 Tornado Cash 的交易量和用户多样性立即下降，但净流入量在数月后恢复甚至超过制裁前水平；"
    "处理 Tornado Cash 交易的区块验证者数量持续萎缩，表明审查抵抗能力是脆弱的（fragile）。"
    "2024 年 11 月，美国第五巡回上诉法院裁定 OFAC 制裁不可变智能合约超越了其法定权力；2025 年 3 月，美国财政部正式解除了对 Tornado Cash 的制裁。")

add("Normal",
    "这一系列事件表明：单纯制裁隐私工具不是长久之计，需要从技术层面寻找隐私与合规的平衡点。")

add("Heading 3", "2.4.2 Privacy Pools：隐私与合规的实用均衡")

add("Normal",
    "Buterin, Illum, Nadler, Schär & Soleimani（2023）在论文 Blockchain Privacy and Regulatory Compliance: Towards a Practical Equilibrium"
    "（发表于 Blockchain: Research and Applications）中 [23]，提出了 Privacy Pools 协议和 Association Sets（关联集合）的核心概念："
    "（1）存款阶段：用户将资金存入隐私池（与 Tornado Cash 类似）；"
    "（2）取款阶段：用户选择一个 Association Set，并用零知识证明（ZKP）证明"我的存款属于这个集合"，但不暴露具体是哪一笔；"
    "（3）包含集合（Inclusion）："我的存款属于已知合法来源的存款集合"；"
    "排除集合（Exclusion）："我的存款不属于 OFAC 制裁地址的存款集合"。")

add("Normal",
    "该协议首次在学术上证明了"隐私"与"合规"不必是非此即彼的——"
    "用户可以在保护交易细节的同时，向验证方证明资金来源的合法性。Privacy Pools v1 已于 2024 年在 Ethereum 主网上线。")

add("Heading 3", "2.4.3 Proof of Innocence 的实践与局限")

add("Normal",
    "Privacy Pools 的核心思想催生了 Proof of Innocence（无辜证明）的工程实践。"
    "Railgun 协议率先部署了 Private Proofs of Innocence（PPOI）系统：在用户存入代币（shield）时，"
    "钱包自动生成一个 ZK 证明，证明该代币不属于预设的非法交易/地址列表。"
    "该证明由去中心化的 POI 节点验证，整个过程端到端加密，不暴露用户的地址、余额或交易历史。")

add("Normal",
    "然而，Constantinides & Cartlidge（2025）在 zkMixer: A Configurable Zero-Knowledge Mixer with Anti-Money Laundering Consensus Protocols"
    "（已被 IEEE DAPPS 2025 接收）中 [24]，指出了 Proof of Innocence 在实践中的根本缺陷："
    "PoI 依赖于黑名单的完整性和实时性——如果一笔存款在通过 PoI 检查之后才被标记为非法，"
    "则该存款已经进入隐私池且不可撤销。该论文提出了替代方案：通过共识机制在存款进入混币池之前由参与者集体验证，"
    "若未通过验证，则可冻结或退回存款。")

add("Heading 3", "2.4.4 隐私保护 AML 的密码学技术路线")

add("Normal",
    "在 ZKP 之外，其他密码学方法也被探索用于隐私保护 AML。"
    "Effendi & Chattopadhyay（2024）在 Privacy-Preserving Graph-Based Machine Learning with Fully Homomorphic Encryption "
    "for Collaborative Anti-Money Laundering（SPACE 2024 会议）中 [25]，展示了使用全同态加密（FHE）在加密数据上直接执行图机器学习"
    "（XGBoost 达到 99%+ 准确率），使得多个金融机构可以在不共享原始数据的前提下协作完成 AML 检测。"
    "Chaudhary（2023）提出的 zkFi 框架 [26] 将零知识证明封装为 DeFi 协议的即插即用合规插件，降低了 ZKP 的集成门槛。")

add("Heading 3", "2.4.5 与本研究的关系")

add("Normal",
    "本研究目前基于公开链上数据进行风险追踪，追踪结果（完整的 BFS 路径树）暴露了用户的交易关联信息。"
    "从隐私保护的角度，一个自然的演进方向是：用户在本地运行追踪，生成一个 ZK 证明——"
    ""该地址在 N 跳内没有黑名单关联"——但不暴露具体路径。"
    "验证方（交易所或交易对手）只看到证明有效与否，看不到追踪过程中涉及的任何中间地址。"
    "这一方向面临的主要挑战包括：（1）黑名单的实时更新问题——证明的有效性依赖于黑名单的某个快照版本；"
    "（2）计算开销——对 BFS 追踪树生成 ZK 证明的电路复杂度较高；"
    "（3）Association Set 的构建者信任问题——谁来决定哪些地址属于"合法"集合。"
    "这些是当前学术界的开放问题，超出本研究的当前范围，但作为长期研究方向具有明确的价值。")

# ========== 3. 插入段落到文档 ==========
# python-docx 没有直接的 insert_paragraph 方法
# 需要通过操作底层 XML 来在指定位置插入

from docx.oxml.ns import qn
from lxml import etree

body = doc.element.body
# 获取 insert_idx 对应的 XML 元素
ref_element = doc.paragraphs[insert_idx]._element

for style, text in new_paragraphs:
    # 创建新的段落元素
    new_p = deepcopy(doc.paragraphs[0]._element)  # 复制一个段落作为模板
    # 清除内容
    for child in list(new_p):
        new_p.remove(child)

    # 设置段落属性
    pPr = new_p.makeelement(qn('w:pPr'), {})

    # 根据 style 设置对应的样式
    pStyle = pPr.makeelement(qn('w:pStyle'), {qn('w:val'): style})
    pPr.append(pStyle)
    new_p.append(pPr)

    # 添加文本 run
    run = new_p.makeelement(qn('w:r'), {})
    rPr = run.makeelement(qn('w:rPr'), {})
    run.append(rPr)

    t = run.makeelement(qn('w:t'), {})
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    run.append(t)
    new_p.append(run)

    # 插入到 ref_element 之前
    ref_element.addprevious(new_p)

# ========== 4. 追加新的参考文献 ==========
# 找到最后一条参考文献，在其后追加
last_ref_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith("[16]"):
        last_ref_idx = i
        break

if last_ref_idx is None:
    # fallback: 在文档末尾追加
    print("Warning: 找不到 [16]，在末尾追加参考文献")

new_refs = [
    "[17] European Union (2024). Regulation (EU) 2024/1689 — AI Act, Annex III: High-Risk AI Systems. https://artificialintelligenceact.eu/annex/3/",
    "[18] Kute, D. et al. (2026). Explainable and Fair Anti-Money Laundering Models Using a Reproducible SHAP Framework for Financial Institutions. Discover Artificial Intelligence. Springer.",
    "[19] Watson, A., Richards, G., & Schiff, D. (2025). Explain First, Trust Later: LLM-Augmented Explanations for Graph-Based Crypto Anomaly Detection. arXiv:2506.14933.",
    "[20] Nicholls, J. et al. (2024). Large Language Model XAI Approach for Illicit Activity Investigation in Bitcoin. Neural Computing and Applications. Springer.",
    "[21] Sun, H. et al. (2024). Large Language Models for Blockchain Security: A Systematic Literature Review. arXiv:2403.14280.",
    "[22] Brownworth, A., Durfee, J., Lee, M., & Martin, A. (2024). Regulating Decentralized Systems: Evidence from Sanctions on Tornado Cash. Federal Reserve Bank of New York Staff Reports, No. 1112.",
    "[23] Buterin, V., Illum, J., Nadler, M., Schär, F., & Soleimani, A. (2023). Blockchain Privacy and Regulatory Compliance: Towards a Practical Equilibrium. Blockchain: Research and Applications, 5(1), 100176.",
    "[24] Constantinides, T. & Cartlidge, J. (2025). zkMixer: A Configurable Zero-Knowledge Mixer with Anti-Money Laundering Consensus Protocols. arXiv:2503.14729. IEEE DAPPS 2025.",
    "[25] Effendi, F. & Chattopadhyay, A. (2024). Privacy-Preserving Graph-Based Machine Learning with Fully Homomorphic Encryption for Collaborative Anti-Money Laundering. SPACE 2024. arXiv:2411.02926.",
    "[26] Chaudhary, A. (2023). zkFi: Privacy-Preserving and Regulation Compliant Transactions using Zero Knowledge Proofs. arXiv:2307.00521.",
]

if last_ref_idx is not None:
    ref_after = doc.paragraphs[last_ref_idx]._element
    for ref_text in reversed(new_refs):
        new_p = deepcopy(doc.paragraphs[last_ref_idx]._element)
        # 清除内容
        for child in list(new_p):
            if child.tag != qn('w:pPr'):
                new_p.remove(child)
        # 添加文本
        run = new_p.makeelement(qn('w:r'), {})
        rPr = run.makeelement(qn('w:rPr'), {})
        run.append(rPr)
        t = run.makeelement(qn('w:t'), {})
        t.text = ref_text
        t.set(qn('xml:space'), 'preserve')
        run.append(t)
        new_p.append(run)
        ref_after.addnext(new_p)

# ========== 5. 保存 ==========
doc.save(OUTPUT)
print(f"已保存到: {OUTPUT}")
